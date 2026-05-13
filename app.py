import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
from PIL import Image
import pandas as pd
import os
import folium
from folium.plugins import LocateControl
from streamlit_folium import st_folium
from fpdf import FPDF

# 1. Configuração inicial do Streamlit
st.set_page_config(page_title="DF LEGAL - Vistoria ERB", layout="wide")

if 'lista_autos' not in st.session_state:
    st.session_state.lista_autos = []

# Inicialização das coordenadas
if 'lat' not in st.session_state:
    st.session_state.lat = -15.8262933
if 'long' not in st.session_state:
    st.session_state.long = -47.815471

# 2. Localização Lateral (INTERATIVA E MAIOR)
st.sidebar.header("📍 Localização")
st.sidebar.write("Escolha a visão (Rua/Satélite) no ícone do mapa:")

col_lat, col_long = st.sidebar.columns(2)
lat_input = col_lat.number_input("Lat", value=st.session_state.lat, format="%.7f")
long_input = col_long.number_input("Long", value=st.session_state.long, format="%.7f")

st.session_state.lat = lat_input
st.session_state.long = long_input

m = folium.Map(location=[st.session_state.lat, st.session_state.long], zoom_start=16)

# Camadas de mapa
folium.TileLayer('openstreetmap', name='Mapa de Ruas').add_to(m)
folium.TileLayer(
    tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
    attr='Esri', name='Satélite', overlay=False, control=True
).add_to(m)
folium.LayerControl().add_to(m)

LocateControl(position="topleft").add_to(m)
folium.Marker([st.session_state.lat, st.session_state.long], icon=folium.Icon(color="red", icon="info-sign")).add_to(m)

output = st_folium(m, height=450, use_container_width=True, key="mapa_sidebar")

if output and output.get("last_clicked"):
    click_lat = output["last_clicked"]["lat"]
    click_lng = output["last_clicked"]["lng"]
    if click_lat != st.session_state.lat or click_lng != st.session_state.long:
        st.session_state.lat = click_lat
        st.session_state.long = click_lng
        st.rerun()

# 3. Cabeçalho Visual no App
st.markdown("""
    <div style="text-align: center;">
        <h2 style="margin-bottom: 0;">GOVERNO DO DISTRITO FEDERAL</h2>
        <p style="margin-top: 0;">Secretaria de Estado de Proteção da Ordem Urbanística do Distrito Federal - DF LEGAL<br>
        Subsecretaria de Fiscalização de Obras</p>
    </div>
    <hr>
""", unsafe_allow_html=True)

# 4. Formulário de Dados
with st.expander("📝 Dados da Ação Fiscal", expanded=True):
    c1, c2 = st.columns(2)
    with c1:
        processo = st.text_input("Número do processo SEI")
        empresa = st.text_input("Empresa responsável")
        cnpj = st.text_input("CNPJ")
        regiao = st.text_input("Região Administrativa")
    with c2:
        endereco = st.text_input("Endereço Completo")
        st.text_input("Coordenadas (Lat, Long)", value=f"{st.session_state.lat}, {st.session_state.long}", disabled=True)
        licenca = st.text_input("Número da licença")
        certificado = st.text_input("Número do Certificado")

# 5. Triagem e Fluxo de Decisão
st.header("🚦 Triagem da Vistoria")
licenciada = st.selectbox("1 – A Estação Rádio Base foi licenciada?", ["Sim", "Não", "Dispensa de licenciamento"])

ir_para_checklist = False
texto_relato = ""
texto_conclusao = ""

if licenciada == "Não":
    texto_relato = "Na vistoria realizada, verificou-se a existência de implantação de infraestrutura de telecomunicações sem o devido licenciamento."
    texto_conclusao = "Tendo em vista a ausência de licenciamento, foram realizadas as ações fiscais indicadas no item 'Autos Fiscais', visando a regularização da infraestrutura implantada."
elif licenciada in ["Sim", "Dispensa de licenciamento"]:
    implantada = st.radio("2 - A ERB está implantada?", ["Sim", "Não"])
    if implantada == "Sim":
        ir_para_checklist = True
    else:
        texto_relato = "Na vistoria realizada, verificou-se que a infraestrutura de telecomunicações, devidamente licenciada, ainda não foi implantada."
        texto_conclusao = "Tendo em vista que a infraestrutura de telecomunicações ainda não foi implantada, informa-se que, por ora, não há necessidade de adoção de medidas fiscais adicionais, ressalvando-se que o presente processo será oportunamente incluído em programação fiscal."

if ir_para_checklist:
    st.header("📝 Checklist de Vistoria")
    col_q1, col_q2 = st.columns(2)
    op = ["Sim", "Não", "Não se aplica"]
    with col_q1:
        r1 = st.radio("A área pública foi recuperada?", op)
        r2 = st.radio("A altura da ERB está de acordo com o projeto?", op)
    with col_q2:
        r3 = st.radio("O cercamento foi executado conforme o projeto?", op)
        r4 = st.radio("A placa de advertência foi instalada?", op)
    
    irr_list = []
    if r1 == "Não": irr_list.append("- Falta recuperar a área pública degradada em decorrência da obra")
    if r2 == "Não": irr_list.append("- A altura da infraestrutura de telecomunicação implantada está em desacordo com o projeto aprovado. Em caso de discordância, apresentar o laudo topográfico.")
    if r3 == "Não": irr_list.append("- Não foi realizado o cercamento conforme o projeto aprovado")
    if r4 == "Não": irr_list.append("- Falta a placa de advertência conforme o projeto aprovado")

    if not irr_list:
        texto_relato = "Na vistoria realizada, verificou-se a inexistência de irregularidades, estando o equipamento em conformidade com o projeto aprovado."
        texto_conclusao = "Constatou-se que o equipamento encontra-se implantado em conformidade com o projeto aprovado, não tendo, até o presente momento, a necessidade de adoção de medidas fiscais adicionais."
    else:
        texto_relato = "Na vistoria realizada, foram constatadas as seguintes irregularidades:\n" + "\n".join(irr_list)
        texto_conclusao = "Tendo em vista as irregularidades listadas no item 'Relato da Vistoria', foram realizadas as ações fiscais indicadas no item 'Autos Fiscais', visando a regularização da Estação Rádio Base."

# --- NOVO CAMPO: OBSERVAÇÕES ---
st.header("💬 Observações Adicionais")
observacoes_input = st.text_area("Digite observações relevantes (opcional)", placeholder="Caso não preenchido, não aparecerá no relatório.")

# 6. Autos Fiscais
st.header("⚖️ Autos Fiscais")
with st.container(border=True):
    ac1, ac2, ac3 = st.columns([2, 2, 1])
    t_auto = ac1.selectbox("Tipo", ["Auto de Notificação", "Auto de Infração", "Intimação Demolitória"])
    n_auto = ac2.text_input("Número do Auto")
    if ac3.button("➕ Adicionar"):
        if n_auto:
            st.session_state.lista_autos.append({"tipo": t_auto, "numero": n_auto})
            st.rerun()

for i, a in enumerate(st.session_state.lista_autos):
    st.write(f"✅ {a['tipo']} nº {a['numero']}")

# 7. Registro Fotográfico
st.header("📸 Registro Fotográfico")
fotos_upload = st.file_uploader("Upload de fotos", accept_multiple_files=True, type=['jpg', 'png', 'jpeg'])

# 8. Função de Geração de Documento
def gerar_docx(dados, relato, autos_lista, conclusao, observacoes, fotos):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    def force_arial(run):
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # Margens
    section = doc.sections[0]
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(5.5) 
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)
    
    # =========================================================
    # --- CABEÇALHO OFICIAL BASEADO NA RÉGUA ---
    # =========================================================
    header = section.header
    
    # Limpa parágrafos vazios do cabeçalho
    for paragraph in header.paragraphs:
        p_elem = paragraph._element
        p_elem.getparent().remove(p_elem)
        
    htable = header.add_table(1, 3, width=Cm(19.0))
    htable.style = 'Table Grid' # Bordas ativas
    
    # Largura das colunas baseadas na régua
    for cell in htable.columns[0].cells: cell.width = Cm(4.0)
    for cell in htable.columns[1].cells: cell.width = Cm(11.0)
    for cell in htable.columns[2].cells: cell.width = Cm(4.0)
    
    # Centralização vertical da tabela inteira
    for cell in htable.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Coluna 1: Logo GDF (Centralizada dentro dos 4cm)
    p_left = htable.cell(0, 0).paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists("logo_gdf.png"):
        p_left.add_run().add_picture("logo_gdf.png", width=Cm(2.2), height=Cm(3.09))
    
    # Coluna 2: Texto Central (Nos 11cm de espaço)
    p_center = htable.cell(0, 1).paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_center.paragraph_format.space_before = Pt(0)
    p_center.paragraph_format.space_after = Pt(0)
    
    run1 = p_center.add_run("GOVERNO DO DISTRITO FEDERAL\n")
    run1.bold = True
    run1.font.name = 'Arial'
    run1.font.size = Pt(11)
    
    run2 = p_center.add_run("Secretaria de Estado de Proteção da Ordem Urbanística do\nDistrito Federal - DF LEGAL\nSubsecretaria de Fiscalização de Obras")
    run2.bold = False
    run2.font.name = 'Arial'
    run2.font.size = Pt(11)
    
    # Coluna 3: Logo DF Legal (Centralizada dentro dos 4cm)
    p_right = htable.cell(0, 2).paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists("logo_dflegal.png"):
        p_right.add_run().add_picture("logo_dflegal.png", width=Cm(2.28), height=Cm(2.94))

    # Corpo do Documento
    doc.add_paragraph() 
    
    def add_campo(label, valor):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE 
        p.paragraph_format.space_before = Pt(0) 
        p.paragraph_format.space_after = Pt(0)  
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        force_arial(rl)
        rv = p.add_run(str(valor))
        force_arial(rv)

    add_campo("PROCESSO SEI", dados['processo'])
    add_campo("EMPRESA", dados['empresa'])
    add_campo("CNPJ", dados['cnpj'])
    add_campo("REGIÃO ADMINISTRATIVA", dados['regiao'])
    add_campo("ENDEREÇO", dados['endereco'])
    add_campo("COORDENADAS", f"{dados['lat']}, {dados['long']}")
    if dados['licenca']: add_campo("LICENÇA", dados['licenca'])
    if dados['certificado']: add_campo("CERTIFICADO", dados['certificado'])
    
    def add_secao(titulo, conteudo):
        doc.add_paragraph()
        p_t = doc.add_paragraph()
        rt = p_t.add_run(titulo)
        rt.bold = True
        force_arial(rt)
        p_t.paragraph_format.line_spacing = 1.5
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_c.paragraph_format.line_spacing = 1.5
        p_c.paragraph_format.first_line_indent = Cm(1.25)
        rc = p_c.add_run(conteudo)
        force_arial(rc)

    intro = "Trata-se de vistoria de implantação de infraestrutura de telecomunicações, nos termos da Lei Complementar nº 971/2020, regulamentada pelo Decreto nº 41446/2020. A vistoria consistiu basicamente em verificar a compatibilidade entre a infraestrutura implantada e o projeto aprovado, sem adentrar nos aspectos técnicos construtivos referentes à infraestrutura."
    add_secao("1. INTRODUÇÃO", intro)
    add_secao("2. RELATO DA VISTORIA", relato)
    
    # Seção 3: Autos Fiscais
    doc.add_paragraph()
    p_t_autos = doc.add_paragraph()
    rta = p_t_autos.add_run("3. AUTOS FISCAIS")
    rta.bold = True
    force_arial(rta)
    p_a = doc.add_paragraph()
    p_a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_a.paragraph_format.line_spacing = 1.5
    p_a.paragraph_format.first_line_indent = Cm(1.25)
    if not autos_lista:
        ra = p_a.add_run("Não houve necessidade de emissão de autos fiscais.")
        force_arial(ra)
    else:
        texto_autos = "\n".join([f"- {a['tipo']} nº {a['numero']}" for a in autos_lista])
        ra = p_a.add_run(f"Foram emitidos os seguintes autos:\n{texto_autos}")
        force_arial(ra)

    add_secao("4. CONCLUSÃO", conclusao)

    # --- LÓGICA CONDICIONAL: OBSERVAÇÕES ---
    proxima_secao_num = 5
    if observacoes and observacoes.strip():
        add_secao(f"{proxima_secao_num}. OBSERVAÇÕES", observacoes)
        proxima_secao_num += 1

    if fotos:
        doc.add_paragraph()
        p_tf = doc.add_paragraph()
        rtf = p_tf.add_run(f"{proxima_secao_num}. REGISTRO FOTOGRÁFICO")
        rtf.bold = True
        force_arial(rtf)
        p_tf.paragraph_format.line_spacing = 1.5
        for f in fotos:
            img = Image.open(f)
            if img.mode != 'RGB': img = img.convert('RGB')
            img.thumbnail((600, 600))
            tmp = BytesIO()
            img.save(tmp, format="JPEG")
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(tmp, width=Cm(12))

    footer = section.footer
    fpara = footer.paragraphs[0]
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fpara.add_run("SIA trecho 3 Lotes 1545 e 1555 | CEP: 71.200-039 | Telefone: (61) 3961-5125\nwww.dflegal.df.gov.br")
    fr.font.size = Pt(8)
    force_arial(fr)

    bio = BytesIO(); doc.save(bio)
    return bio.getvalue()

# 9. Botão Final
st.divider()
if st.button("🚀 PREPARAR RELATÓRIOS"):
    # MONTANDO O DICIONÁRIO COMPLETO PARA EVITAR KEYERROR
    dados_f = {
        "processo": processo,
        "empresa": empresa,
        "cnpj": cnpj,              # Adicionado
        "regiao": regiao,          # Adicionado
        "endereco": endereco,
        "lat": st.session_state.lat,
        "long": st.session_state.long,
        "licenca": licenca,        # Adicionado
        "certificado": certificado # Adicionado
    }
    
    # Gerando os arquivos
    docx_bytes = gerar_docx(dados_f, texto_relato, st.session_state.lista_autos, texto_conclusao, observacoes_input, fotos_upload)
    
    # Nota: Certifique-se de que a função gerar_pdf esteja definida no seu código 
    # ou use a lógica de download apenas para o Word se ainda não a implementou totalmente.
    try:
        pdf_bytes = gerar_pdf(dados_f, texto_relato, st.session_state.lista_autos, texto_conclusao, observacoes_input, fotos_upload)
        col_d1, col_d2 = st.columns(2)
        col_d1.download_button("⬇️ Baixar Word", data=docx_bytes, file_name=f"Vistoria_{processo}.docx")
        col_d2.download_button("⬇️ Baixar PDF", data=pdf_bytes, file_name=f"Vistoria_{processo}.pdf")
    except NameError:
        # Caso a função gerar_pdf não exista, oferece apenas o Word
        st.download_button("⬇️ Baixar Word", data=docx_bytes, file_name=f"Vistoria_{processo}.docx")